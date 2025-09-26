import tkinter as tk
from tkinter import ttk, messagebox, filedialog, scrolledtext
import sqlite3
import json
import os
import logging
import threading
from datetime import datetime
import webbrowser
import re
from collections import Counter
import shutil
from pathlib import Path
import sys
from contextlib import contextmanager
from dataclasses import dataclass
from typing import Optional, List, Dict, Any, Tuple, Generator
from concurrent.futures import ThreadPoolExecutor, Future

# Import AI APIs
from ai_apis import AIManager

# Import additional libraries
try:
    from wordcloud import WordCloud
    WORDCLOUD_AVAILABLE = True
except ImportError:
    WORDCLOUD_AVAILABLE = False
    
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


@dataclass
class TextStats:
    """Dataclass to hold text statistics."""
    char_count: int = 0
    word_count: int = 0
    sentence_count: int = 0
    line_count: int = 0
    token_count: int = 0


class SettingsManager:
    """Encapsulates settings access with validation and a single save method."""
    def __init__(self, settings_file: str = 'settings.json'):
        self.filepath = settings_file
        self.settings: Dict[str, Any] = {}
        self.load()

    def _get_defaults(self) -> Dict[str, Any]:
        """Provides default settings."""
        return {
            'export_path': str(Path.home() / 'Downloads'),
            'ai_provider': 'OpenAI',
            'ai_api_key': '',
            'log_level': 'INFO',
            'window_geometry': '1200x800+100+100'
        }

    def load(self) -> None:
        """Loads settings from the settings file, applying defaults for missing keys."""
        defaults = self._get_defaults()
        try:
            if os.path.exists(self.filepath):
                with open(self.filepath, 'r') as f:
                    loaded_settings = json.load(f)
                # Merge loaded settings with defaults to ensure all keys exist
                self.settings = {**defaults, **loaded_settings}
            else:
                self.settings = defaults
                self.save()
        except (json.JSONDecodeError, IOError) as e:
            # If loading fails, fall back to defaults
            print(f"Error loading settings: {e}. Using default settings.")
            self.settings = defaults

    def save(self) -> None:
        """Saves the current settings to the settings file."""
        try:
            with open(self.filepath, 'w') as f:
                json.dump(self.settings, f, indent=2)
        except IOError as e:
            print(f"Error saving settings: {e}")

    def get(self, key: str, default: Optional[Any] = None) -> Any:
        """Gets a setting value by key."""
        return self.settings.get(key, default)

    def set(self, key: str, value: Any) -> None:
        """Sets a setting value and immediately persists it."""
        self.settings[key] = value
        self.save()


class PromptMiniApp:
    def __init__(self) -> None:
        self.root = tk.Tk()
        self.root.title("Prompt Mini")
        
        self.settings_manager = SettingsManager()
        
        window_geometry = self.settings_manager.get('window_geometry', '1200x800+100+100')
        self.root.geometry(window_geometry)
        
        self.setup_logging()
        self.apply_log_level()
        self.init_database()
        
        self.search_debounce_timer: Optional[str] = None
        self.text_debounce_timer: Optional[str] = None
        
        self.selected_items: List[int] = []
        self.current_item: Optional[int] = None
        
        self.editing_mode: bool = False
        
        self.sort_column: Optional[str] = None
        self.sort_direction: Optional[str] = None
        
        self.prompt_cache: Dict[int, Tuple] = {}
        self.logger.info("Initialized prompt cache")
        
        # Thread pool for cancellable searches
        self.search_executor = ThreadPoolExecutor(max_workers=1)
        self.current_search_future: Optional[Future] = None
        
        self.create_menu()
        self.create_main_ui()
        
        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)
        
        self.perform_search(select_first=True)
        
    def setup_logging(self) -> None:
        """Set up the application's logging system."""
        self.log_handler = logging.StreamHandler()
        self.log_handler.setLevel(logging.DEBUG)
        formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
        self.log_handler.setFormatter(formatter)
        
        self.logger = logging.getLogger('PromptMini')
        self.logger.setLevel(logging.DEBUG)
        self.logger.addHandler(self.log_handler)
        
        self.log_messages: List[Tuple[int, str]] = []
        
        class LogCapture(logging.Handler):
            def __init__(self, app: 'PromptMiniApp'):
                super().__init__()
                self.app = app
                
            def emit(self, record: logging.LogRecord) -> None:
                msg = self.format(record)
                self.app.log_messages.append((record.levelno, msg))
                if len(self.app.log_messages) > 1000:
                    self.app.log_messages = self.app.log_messages[-500:]
                
                if hasattr(self.app, 'status_bar'):
                    parts = msg.split(' - ')
                    if len(parts) >= 3:
                        status_msg = parts[-1].strip()
                        self.app.update_status_bar(status_msg)
        
        self.log_capture = LogCapture(self)
        self.log_capture.setFormatter(formatter)
        self.logger.addHandler(self.log_capture)
        
    def apply_log_level(self) -> None:
        """Apply the log level from settings."""
        level_map = {
            "DEBUG": logging.DEBUG, "INFO": logging.INFO, "WARNING": logging.WARNING,
            "ERROR": logging.ERROR, "CRITICAL": logging.CRITICAL
        }
        log_level_str = self.settings_manager.get('log_level', 'INFO')
        level = level_map.get(log_level_str, logging.INFO)
        self.logger.setLevel(level)
        self.log_handler.setLevel(level)
            
    def auto_size_window(self, window: tk.Toplevel, min_width: int = 800, min_height: int = 600, show_window: bool = True) -> None:
        """Automatically size a window to fit content with minimum dimensions."""
        window.update_idletasks()
        
        screen_width = window.winfo_screenwidth()
        screen_height = window.winfo_screenheight()
        
        req_width = max(min_width, window.winfo_reqwidth() + 50)
        req_height = max(min_height, window.winfo_reqheight() + 100)
        
        max_width = int(screen_width * 0.9)
        max_height = int(screen_height * 0.9)
        
        final_width = min(req_width, max_width)
        final_height = min(req_height, max_height)
        
        x = (screen_width - final_width) // 2
        y = (screen_height - final_height) // 2
        
        window.geometry(f"{final_width}x{final_height}+{x}+{y}")
        
        if show_window:
            window.deiconify()
            
    @contextmanager
    def get_db_connection(self) -> Generator[sqlite3.Connection, None, None]:
        """Provide a managed database connection."""
        conn = None
        try:
            conn = sqlite3.connect('prompt_mini.db', timeout=10.0)
            conn.row_factory = sqlite3.Row
            conn.execute('PRAGMA foreign_keys = ON')
            conn.execute('PRAGMA journal_mode=WAL')
            yield conn
        except Exception as e:
            self.logger.error(f"Database connection error: {e}")
            if conn:
                conn.rollback()
            raise
        finally:
            if conn:
                conn.close()

    def init_database(self) -> None:
        """Initialize the SQLite database and Full-Text Search (FTS5) table."""
        try:
            with self.get_db_connection() as conn:
                conn.execute('''
                    CREATE TABLE IF NOT EXISTS prompts (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        Created DATETIME DEFAULT CURRENT_TIMESTAMP,
                        Modified DATETIME DEFAULT CURRENT_TIMESTAMP,
                        Purpose TEXT(255),
                        Prompt TEXT,
                        SessionURLs TEXT,
                        Tags TEXT,
                        Note TEXT
                    )
                ''')
                
                # Drop legacy triggers and FTS table to ensure schema is correct
                for trigger in ['prompts_after_insert', 'prompts_after_delete', 'prompts_after_update', 'prompts_ai', 'prompts_ad', 'prompts_au']:
                    conn.execute(f'DROP TRIGGER IF EXISTS {trigger}')
                conn.execute('DROP TABLE IF EXISTS prompts_fts')
                
                conn.execute('''
                    CREATE VIRTUAL TABLE prompts_fts USING fts5(
                        Purpose, Prompt, SessionURLs, Tags, Note,
                        content='prompts',
                        content_rowid='id'
                    )
                ''')
                
                conn.executescript('''
                    CREATE TRIGGER IF NOT EXISTS prompts_after_insert AFTER INSERT ON prompts BEGIN
                        INSERT INTO prompts_fts(rowid, Purpose, Prompt, SessionURLs, Tags, Note)
                        VALUES (new.id, new.Purpose, new.Prompt, new.SessionURLs, new.Tags, new.Note);
                    END;
                    CREATE TRIGGER IF NOT EXISTS prompts_after_delete AFTER DELETE ON prompts BEGIN
                        INSERT INTO prompts_fts(prompts_fts, rowid, Purpose, Prompt, SessionURLs, Tags, Note)
                        VALUES ('delete', old.id, old.Purpose, old.Prompt, old.SessionURLs, old.Tags, old.Note);
                    END;
                    CREATE TRIGGER IF NOT EXISTS prompts_after_update AFTER UPDATE ON prompts BEGIN
                        INSERT INTO prompts_fts(prompts_fts, rowid, Purpose, Prompt, SessionURLs, Tags, Note)
                        VALUES ('delete', old.id, old.Purpose, old.Prompt, old.SessionURLs, old.Tags, old.Note);
                        INSERT INTO prompts_fts(rowid, Purpose, Prompt, SessionURLs, Tags, Note)
                        VALUES (new.id, new.Purpose, new.Prompt, new.SessionURLs, new.Tags, new.Note);
                    END;
                ''')
                
                conn.execute('INSERT INTO prompts_fts(prompts_fts) VALUES("rebuild")')
                conn.commit()
            self.logger.info("Database initialized successfully")
        except Exception as e:
            self.logger.error(f"Database initialization error: {e}")
            messagebox.showerror("Database Error", f"Failed to initialize database: {e}")    
        
    def create_menu(self) -> None:
        """Create the main application menu bar."""
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)
        
        file_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="File", menu=file_menu)
        file_menu.add_command(label="Export Location", command=self.set_export_location)
        file_menu.add_separator()
        
        export_view_menu = tk.Menu(file_menu, tearoff=0)
        file_menu.add_cascade(label="Export View", menu=export_view_menu)
        for fmt in ['csv', 'pdf', 'txt', 'docx']:
            export_view_menu.add_command(label=fmt.upper(), command=lambda f=fmt: self.export_view(f))

        export_all_menu = tk.Menu(file_menu, tearoff=0)
        file_menu.add_cascade(label="Export All", menu=export_all_menu)
        for fmt in ['csv', 'pdf', 'txt', 'docx']:
            export_all_menu.add_command(label=fmt.upper(), command=lambda f=fmt: self.export_all(f))
        
        file_menu.add_separator()
        file_menu.add_command(label="Backup", command=self.backup_database)
        file_menu.add_command(label="Restore", command=self.restore_database)
        file_menu.add_command(label="Import", command=self.import_database)
        file_menu.add_separator()
        file_menu.add_command(label="Exit", command=self.root.quit)
        
        settings_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Settings", menu=settings_menu)
        settings_menu.add_command(label="Console Log", command=self.show_console_log)
        
        help_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Help", menu=help_menu)
        help_menu.add_command(label="GitHub", command=lambda: webbrowser.open("https://github.com/matbanik/prompt_mini"))
        
    def create_main_ui(self) -> None:
        """Create the main user interface components."""
        search_frame = ttk.Frame(self.root)
        search_frame.pack(fill=tk.X, padx=10, pady=5)
        
        help_btn = ttk.Button(search_frame, text="?", width=3, command=self.show_search_help)
        help_btn.pack(side=tk.LEFT, padx=(0, 5))
        
        self.search_var = tk.StringVar()
        self.search_var.trace_add('write', self.on_search_change)
        self.search_entry = ttk.Entry(search_frame, textvariable=self.search_var, width=80)
        self.search_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        
        search_btn = ttk.Button(search_frame, text="Search", command=lambda: self.perform_search())
        search_btn.pack(side=tk.RIGHT, padx=(5, 0))
        
        main_frame = ttk.Frame(self.root)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        try:
            self.paned_window = ttk.PanedWindow(main_frame, orient=tk.HORIZONTAL, sashwidth=8, sashrelief=tk.RAISED)
        except tk.TclError:
            self.paned_window = ttk.PanedWindow(main_frame, orient=tk.HORIZONTAL)
        self.paned_window.pack(fill=tk.BOTH, expand=True)
        
        left_frame = ttk.Frame(self.paned_window)
        self.paned_window.add(left_frame, weight=7)
        
        columns = ('ID', 'Created', 'Modified', 'Purpose', 'Tags')
        self.tree = ttk.Treeview(left_frame, columns=columns, show='headings', selectmode='extended')
        
        for col in columns:
            self.tree.heading(col, text=col, command=lambda c=col: self.sort_by_column(c))
        
        self.tree.column('ID', width=40, minwidth=30, stretch=False)
        self.tree.column('Created', width=120, minwidth=100, stretch=False)
        self.tree.column('Modified', width=120, minwidth=100, stretch=False)
        self.tree.column('Purpose', width=200, minwidth=120)
        self.tree.column('Tags', width=150, minwidth=100)
        
        tree_scroll = ttk.Scrollbar(left_frame, orient=tk.VERTICAL, command=self.tree.yview)
        self.tree.configure(yscrollcommand=tree_scroll.set)
        
        self.tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        tree_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.tree.bind('<<TreeviewSelect>>', self.on_tree_select)
        self.tree.bind('<Double-1>', self.on_tree_double_click)
        self.tree.bind('<Motion>', self.on_tree_motion)
        self.tree.bind('<Leave>', self.on_tree_leave)
        
        self.root.bind('<Control-z>', self.undo_text)
        self.root.bind('<Control-y>', self.redo_text)
        
        self.tooltip: Optional[tk.Toplevel] = None
        
        right_frame = ttk.Frame(self.paned_window)
        self.paned_window.add(right_frame, weight=3)
        
        self.btn_frame = ttk.Frame(right_frame)
        self.btn_frame.pack(fill=tk.X, pady=(0, 10))

        self.new_btn = ttk.Button(self.btn_frame, text="New Prompt", command=self.new_item)
        self.duplicate_btn = ttk.Button(self.btn_frame, text="Duplicate", command=self.duplicate_item)
        self.change_btn = ttk.Button(self.btn_frame, text="Change", command=self.change_item)
        self.delete_btn = ttk.Button(self.btn_frame, text="Delete", command=self.delete_items)

        self.save_btn = ttk.Button(self.btn_frame, text="Save", command=self.save_edits)
        self.cancel_btn = ttk.Button(self.btn_frame, text="Cancel", command=self.cancel_edits)

        self.update_action_buttons()
        
        display_frame = ttk.Frame(right_frame)
        display_frame.pack(fill=tk.BOTH, expand=True)
        
        self.create_item_display(display_frame)
        
        self.status_bar = ttk.Label(self.root, text="Ready", relief=tk.SUNKEN, anchor=tk.W)
        self.status_bar.pack(side=tk.BOTTOM, fill=tk.X, padx=5, pady=2)
        
    def update_status_bar(self, message: str) -> None:
        """Update the status bar with a temporary message."""
        if hasattr(self, 'status_bar'):
            self.status_bar.config(text=message)
            self.root.after(5000, lambda: self.status_bar.config(text="Ready"))
            
    def sync_scroll(self, scrollbar: ttk.Scrollbar, line_numbers: tk.Text, *args: str) -> None:
        """Synchronize scrolling between a text widget and its line numbers."""
        scrollbar.set(*args)
        if len(args) >= 2:
            top = float(args[0])
            line_numbers.yview_moveto(top)
            
    def sync_scroll_command(self, main_text: tk.Text, line_numbers: tk.Text, *args: str) -> None:
        """Handle scrollbar commands to sync two text widgets."""
        main_text.yview(*args)
        line_numbers.yview(*args)
        
    def create_item_display(self, parent: ttk.Frame) -> None:
        """Create the widgets for displaying a single prompt item."""
        date_frame = ttk.Frame(parent)
        date_frame.pack(fill=tk.X, pady=(0, 5))
        self.created_label = ttk.Label(date_frame, text="Created: ", foreground="green")
        self.created_label.pack(side=tk.LEFT)
        self.modified_label = ttk.Label(date_frame, text="Modified: ", foreground="blue")
        self.modified_label.pack(side=tk.RIGHT)
        
        self.purpose_frame = ttk.Frame(parent)
        self.purpose_frame.pack(fill=tk.X, pady=(0, 5))
        ttk.Label(self.purpose_frame, text="Purpose:").pack(side=tk.LEFT)
        self.purpose_display = ttk.Label(self.purpose_frame, text="", font=('TkDefaultFont', 9, 'bold'))
        self.purpose_display.pack(side=tk.LEFT, padx=(5, 0))
        
        prompt_frame = ttk.Frame(parent)
        prompt_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 5))
        
        line_frame = ttk.Frame(prompt_frame)
        line_frame.pack(side=tk.LEFT, fill=tk.Y)
        self.line_numbers = tk.Text(line_frame, width=4, padx=3, takefocus=0, border=0, state='disabled', wrap='none')
        self.line_numbers.pack(fill=tk.Y, expand=True)
        
        text_frame = ttk.Frame(prompt_frame)
        text_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)
        self.prompt_display = tk.Text(text_frame, wrap=tk.WORD, state='disabled', undo=True, maxundo=50)
        prompt_scrollbar = ttk.Scrollbar(text_frame, orient=tk.VERTICAL)
        self.prompt_display.config(yscrollcommand=lambda *args: self.sync_scroll(prompt_scrollbar, self.line_numbers, *args))
        prompt_scrollbar.config(command=lambda *args: self.sync_scroll_command(self.prompt_display, self.line_numbers, *args))
        self.prompt_display.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        prompt_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        status_frame = ttk.Frame(parent)
        status_frame.pack(fill=tk.X, pady=(0, 5))
        self.status_label = ttk.Label(status_frame, text="Char: 0 | Word: 0 | Sentence: 0 | Line: 0 | Tokens: 0")
        self.status_label.pack(side=tk.LEFT)
        
        copy_btn = ttk.Button(status_frame, text="Copy", command=self.copy_to_clipboard)
        copy_btn.pack(side=tk.RIGHT, padx=(5, 0))
        tune_btn = ttk.Button(status_frame, text="Tune with AI", command=self.tune_with_ai)
        tune_btn.pack(side=tk.RIGHT)
        
        ttk.Label(parent, text="Session URLs", font=('TkDefaultFont', 9, 'bold')).pack(anchor=tk.W, pady=(0, 2))
        self.urls_display = scrolledtext.ScrolledText(parent, height=7, state='disabled', undo=True, maxundo=50)
        self.urls_display.pack(fill=tk.X, pady=(0, 5))
        
        ttk.Label(parent, text="Tags", font=('TkDefaultFont', 9, 'bold')).pack(anchor=tk.W, pady=(0, 2))
        self.tags_display = ttk.Frame(parent)
        self.tags_display.pack(fill=tk.X, pady=(0, 5))
        
        ttk.Label(parent, text="Note", font=('TkDefaultFont', 9, 'bold')).pack(anchor=tk.W, pady=(0, 2))
        self.note_display = scrolledtext.ScrolledText(parent, height=0, state='disabled', undo=True, maxundo=50)
        self.note_display.pack(fill=tk.BOTH, expand=True, pady=(0, 5))       

    def on_search_change(self, *args: Any) -> None:
        """Handle search input changes with debouncing."""
        if self.search_debounce_timer:
            self.root.after_cancel(self.search_debounce_timer)
        self.search_debounce_timer = self.root.after(300, lambda: self.perform_search())
        
    def perform_search(self, select_item_id: Optional[int] = None, select_first: bool = False) -> None:
        """Perform a cancellable search using a thread pool."""
        search_term = self.search_var.get().strip()
        
        if self.current_search_future and not self.current_search_future.done():
            self.current_search_future.cancel()

        if hasattr(self, 'status_bar'):
            self.status_bar.config(text="Searching...")
            self.root.config(cursor="wait")
            self.root.update_idletasks()
        
        def search_worker(term: str) -> List[Tuple]:
            try:
                with self.get_db_connection() as conn:
                    if term:
                        cursor = conn.execute('''
                            SELECT p.id, p.Created, p.Modified, p.Purpose, p.Prompt, p.SessionURLs, p.Tags, p.Note
                            FROM prompts p JOIN prompts_fts fts ON p.id = fts.rowid
                            WHERE prompts_fts MATCH ? ORDER BY rank
                        ''', (term + '*',))
                    else:
                        cursor = conn.execute('''
                            SELECT id, Created, Modified, Purpose, Prompt, SessionURLs, Tags, Note
                            FROM prompts ORDER BY Modified DESC
                        ''')
                    return cursor.fetchall()
            except Exception as e:
                self.logger.error(f"Search worker error: {e}")
                return []

        self.current_search_future = self.search_executor.submit(search_worker, search_term)
        self.current_search_future.add_done_callback(
            lambda future: self.root.after(0, lambda: self._handle_search_results(future, select_item_id, select_first))
        )
    
    def _handle_search_results(self, future: Future, select_item_id: Optional[int] = None, select_first: bool = False) -> None:
        """Process search results in the main UI thread."""
        if future.cancelled():
            return

        if hasattr(self, 'status_bar'):
            self.status_bar.config(text="Ready")
            self.root.config(cursor="")
        
        error = future.exception()
        if error:
            self.logger.error(f"Search failed: {error}")
            messagebox.showerror("Search Error", f"Search failed: {error}")
            self.search_results = []
        else:
            self.search_results = future.result()
        
        self.refresh_search_view()
        
        if select_item_id:
            self.root.after_idle(lambda: self._select_item_in_tree(select_item_id))
        elif select_first:
            self.root.after_idle(self._select_first_item_in_tree)
            
    def sort_by_column(self, column: str) -> None:
        """Sort the treeview by a specified column, cycling through directions."""
        if self.sort_column == column:
            if self.sort_direction == 'asc':
                self.sort_direction = 'desc'
            else:
                self.sort_column = None
                self.sort_direction = None
        else:
            self.sort_column = column
            self.sort_direction = 'asc'
        
        self.update_column_headers()
        self.refresh_search_view()
        
    def update_column_headers(self) -> None:
        """Update treeview column headers with sort direction indicators."""
        for col in ['ID', 'Created', 'Modified', 'Purpose', 'Tags']:
            text = col
            if col == self.sort_column:
                if self.sort_direction == 'asc':
                    text += " ↑"
                elif self.sort_direction == 'desc':
                    text += " ↓"
            self.tree.heading(col, text=text)
    
    def refresh_search_view(self) -> None:
        """Refresh the search results treeview, applying sorting if active."""
        self.tree.unbind('<<TreeviewSelect>>')
        try:
            self.tree.delete(*self.tree.get_children())
                
            display_results = getattr(self, 'search_results', [])
            
            if self.sort_column and self.sort_direction and display_results:
                col_index = self.tree['columns'].index(self.sort_column)
                reverse = (self.sort_direction == 'desc')
                
                def sort_key(row):
                    val = row[col_index]
                    if self.sort_column == 'ID':
                        return int(val) if val else 0
                    return (val or '').lower()

                display_results = sorted(display_results, key=sort_key, reverse=reverse)
                
            if display_results:
                for row in display_results:
                    tags = row['Tags']
                    tags_display = ""
                    if tags:
                        try:
                            # Strip whitespace and handle potential JSON errors
                            tag_list = [t.strip() for t in (json.loads(tags) if tags.strip().startswith('[') else tags.split(','))]
                            tag_list = [t for t in tag_list if t]
                            tags_display = ', '.join(tag_list[:3])
                            if len(tag_list) > 3:
                                tags_display += "..."
                        except json.JSONDecodeError:
                            self.logger.warning(f"Malformed tags JSON for item {row['id']}: {tags}")
                            tags_display = tags[:30].strip() + "..." if len(tags) > 30 else tags.strip()
                    
                    self.tree.insert('', 'end', values=(
                        row['id'],
                        self.format_datetime(row['Created']),
                        self.format_datetime(row['Modified']),
                        (row['Purpose'] or '')[:50] + "...",
                        tags_display
                    ))
        finally:
            self.tree.bind('<<TreeviewSelect>>', self.on_tree_select)
                
    def format_datetime(self, dt_str: Optional[str]) -> str:
        """Format a datetime string for display."""
        if not dt_str:
            return ""
        try:
            # Handle potential 'Z' suffix for UTC
            dt_str = dt_str.replace('Z', '+00:00')
            dt = datetime.fromisoformat(dt_str)
            return dt.strftime('%Y-%m-%d %I:%M %p')
        except (ValueError, TypeError):
            return dt_str
            
    def on_tree_select(self, event: Optional[tk.Event]) -> None:
        """Handle selection changes in the results treeview."""
        selection = self.tree.selection()
        self.selected_items = [self.tree.item(item)['values'][0] for item in selection]
        
        if len(self.selected_items) == 1:
            self.current_item = self.selected_items[0]
            self.update_item_display()
        else:
            self.current_item = None
            self.clear_item_display()
        
        self.update_action_buttons()

    def on_tree_double_click(self, event: tk.Event) -> None:
        """Handle double-click on a tree item to enter editing mode."""
        if self.current_item:
            self.change_item()
            
    def on_tree_motion(self, event: tk.Event) -> None:
        """Show tooltips for truncated text in the treeview."""
        item = self.tree.identify_row(event.y)
        column = self.tree.identify_column(event.x)
        
        if item and column:
            col_index = int(column.replace('#', '')) - 1
            col_name = self.tree['columns'][col_index]
            
            if col_name in ['Purpose', 'Tags']:
                item_values = self.tree.item(item)['values']
                if not item_values: return
                
                full_text = self.get_full_text_for_tooltip(item, col_name)
                
                if full_text and len(full_text) > len(str(item_values[col_index])):
                    self.show_tooltip(event.x_root, event.y_root, full_text)
                else:
                    self.hide_tooltip()
            else:
                self.hide_tooltip()
        else:
            self.hide_tooltip()
            
    def on_tree_leave(self, event: tk.Event) -> None:
        """Hide tooltip when the mouse leaves the treeview."""
        self.hide_tooltip()

    def _select_item_in_tree(self, item_id: int) -> None:
        """Select an item in the tree by its ID."""
        for item in self.tree.get_children():
            if str(self.tree.item(item)['values'][0]) == str(item_id):
                self.tree.selection_set(item)
                self.tree.focus(item)
                self.tree.see(item)
                self.on_tree_select(None)
                return
        self._select_first_item_in_tree()

    def _select_first_item_in_tree(self) -> None:
        """Select the first item in the tree."""
        children = self.tree.get_children()
        if children:
            self.tree.selection_set(children[0])
            self.tree.focus(children[0])
            self.tree.see(children[0])
            self.on_tree_select(None)
        
    def get_full_text_for_tooltip(self, item_id_str: str, column_name: str) -> str:
        """Retrieve the full text for a tooltip from the cached search results."""
        try:
            item_id = int(self.tree.item(item_id_str)['values'][0])
            if hasattr(self, 'search_results'):
                for row in self.search_results:
                    if row['id'] == item_id:
                        return row[column_name] or ""
            return ""
        except (ValueError, IndexError, Exception) as e:
            self.logger.error(f"Error getting tooltip text: {e}")
            return ""
            
    def show_tooltip(self, x: int, y: int, text: str) -> None:
        """Show a tooltip window at the specified coordinates."""
        self.hide_tooltip()
        self.tooltip = tk.Toplevel(self.root)
        self.tooltip.wm_overrideredirect(True)
        self.tooltip.wm_geometry(f"+{x+10}+{y+10}")
        
        frame = ttk.Frame(self.tooltip, relief=tk.SOLID, borderwidth=1)
        frame.pack()
        ttk.Label(frame, text=text, background="lightyellow", wraplength=400, justify=tk.LEFT).pack(padx=5, pady=5)
        
    def hide_tooltip(self) -> None:
        """Hide the currently visible tooltip."""
        if self.tooltip:
            self.tooltip.destroy()
            self.tooltip = None
            
    def undo_text(self, event: tk.Event) -> str:
        """Handle Ctrl+Z for undo on the focused text widget."""
        focused = self.root.focus_get()
        if isinstance(focused, tk.Text):
            try: focused.edit_undo()
            except tk.TclError: pass
        return "break"
        
    def redo_text(self, event: tk.Event) -> str:
        """Handle Ctrl+Y for redo on the focused text widget."""
        focused = self.root.focus_get()
        if isinstance(focused, tk.Text):
            try: focused.edit_redo()
            except tk.TclError: pass
        return "break"
            
    def update_item_display(self, force_refresh: bool = False) -> None:
        """Update the item display panel, using a cache for performance."""
        if not self.current_item: return
        
        try:
            row = self.prompt_cache.get(self.current_item) if not force_refresh else None
            if row:
                self.logger.debug(f"Using cached data for item {self.current_item}")
            else:
                with self.get_db_connection() as conn:
                    row = conn.execute('SELECT * FROM prompts WHERE id = ?', (self.current_item,)).fetchone()
                    if not row: return
                    
                    if len(self.prompt_cache) > 50:
                        del self.prompt_cache[next(iter(self.prompt_cache))]
                    
                    self.prompt_cache[self.current_item] = row
                    self.logger.debug(f"Fetched and cached data for item {self.current_item}")
                    
            self.created_label.config(text=f"Created: {self.format_datetime(row['Created'])}")
            self.modified_label.config(text=f"Modified: {self.format_datetime(row['Modified'])}")
            self.purpose_display.config(text=row['Purpose'] or "")
            
            for widget in [self.prompt_display, self.urls_display, self.note_display]:
                widget.config(state='normal')
                widget.delete(1.0, tk.END)

            if row['Prompt']: self.prompt_display.insert(1.0, row['Prompt'])
            if row['SessionURLs']:
                self.urls_display.insert(1.0, row['SessionURLs'])
                self.make_urls_clickable()
            if row['Note']: self.note_display.insert(1.0, row['Note'])
            
            for widget in [self.prompt_display, self.urls_display, self.note_display]:
                widget.config(state='disabled')
            
            self.update_line_numbers(row['Prompt'] or "")
            self.update_status(row['Prompt'] or "")
            self.update_tags_display(row['Tags'])
            
        except Exception as e:
            self.logger.error(f"Error updating item display: {e}")
            
    def clear_item_display(self) -> None:
        """Clear all fields in the item display panel."""
        self.created_label.config(text="Created: ")
        self.modified_label.config(text="Modified: ")
        self.purpose_display.config(text="")
        
        for widget in [self.prompt_display, self.urls_display, self.note_display, self.line_numbers]:
            widget.config(state='normal')
            widget.delete(1.0, tk.END)
            widget.config(state='disabled')
        
        self.status_label.config(text="Char: 0 | Word: 0 | Sentence: 0 | Line: 0 | Tokens: 0")
        
        for widget in self.tags_display.winfo_children():
            widget.destroy()
        
    def update_line_numbers(self, text: str) -> None:
        """Update the line numbers displayed next to the prompt text."""
        self.line_numbers.config(state='normal')
        self.line_numbers.delete(1.0, tk.END)
        if text:
            line_count = text.count('\n') + 1
            line_nums = '\n'.join(map(str, range(1, line_count + 1)))
            self.line_numbers.insert(1.0, line_nums)
        self.line_numbers.config(state='disabled')

    def _get_text_statistics(self, text: str) -> TextStats:
        """Calculate statistics for a given block of text."""
        if not text:
            return TextStats()
        
        char_count = len(text)
        word_count = len(text.split())
        sentence_count = len(re.findall(r'[.!?]+', text))
        line_count = text.count('\n') + 1
        token_count = int(word_count * 1.3)  # Rough estimate
        
        return TextStats(char_count, word_count, sentence_count, line_count, token_count)

    def update_status(self, text: str) -> None:
        """Update the status label with text statistics."""
        stats = self._get_text_statistics(text)
        self.status_label.config(
            text=f"Char: {stats.char_count} | Word: {stats.word_count} | "
                 f"Sentence: {stats.sentence_count} | Line: {stats.line_count} | Tokens: {stats.token_count}"
        )
        
    def update_tags_display(self, tags_str: Optional[str]) -> None:
        """Update the tags display with clickable tag buttons."""
        for widget in self.tags_display.winfo_children():
            widget.destroy()
            
        if not tags_str:
            return
            
        try:
            tags = [t.strip() for t in (json.loads(tags_str) if tags_str.strip().startswith('[') else tags_str.split(','))]
            for tag in filter(None, tags):
                btn = ttk.Button(self.tags_display, text=tag, command=lambda t=tag: self.search_by_tag(t))
                btn.pack(side=tk.LEFT, padx=2, pady=2)
        except json.JSONDecodeError:
            self.logger.warning(f"Malformed tags JSON could not be parsed: {tags_str}")
            # Display as raw text if parsing fails
            ttk.Label(self.tags_display, text=tags_str, font=('TkDefaultFont', 8, 'italic')).pack(side=tk.LEFT)
                
    def search_by_tag(self, tag: str) -> None:
        """Perform a search for a specific tag."""
        self.search_var.set(f'Tags: "{tag}"')
        self.perform_search()
        
    def copy_to_clipboard(self) -> None:
        """Copy the current prompt text to the clipboard."""
        if self.current_item:
            try:
                with self.get_db_connection() as conn:
                    prompt = conn.execute('SELECT Prompt FROM prompts WHERE id = ?', (self.current_item,)).fetchone()
                    if prompt and prompt['Prompt']:
                        self.root.clipboard_clear()
                        self.root.clipboard_append(prompt['Prompt'])
                        self.update_status_bar("Prompt text copied to clipboard")
            except Exception as e:
                self.logger.error(f"Copy error: {e}")
                self.update_status_bar(f"Copy failed: {e}")
                
    def make_urls_clickable(self) -> None:
        """Find and tag URLs in the URLs display to make them clickable."""
        content = self.urls_display.get(1.0, tk.END)
        url_pattern = r'https?://[^\s\n]+'
        
        # Remove all existing URL tags
        for tag in self.urls_display.tag_names():
            if tag.startswith("url_"):
                self.urls_display.tag_delete(tag)
        
        for i, line in enumerate(content.splitlines(), 1):
            for match in re.finditer(url_pattern, line):
                start, end = f"{i}.{match.start()}", f"{i}.{match.end()}"
                tag_name = f"url_{i}_{match.start()}"
                self.urls_display.tag_add(tag_name, start, end)
                self.urls_display.tag_config(tag_name, foreground="blue", underline=True)
                
                self.urls_display.tag_bind(tag_name, "<Enter>", lambda e: self.urls_display.config(cursor="hand2"))
                self.urls_display.tag_bind(tag_name, "<Leave>", lambda e: self.urls_display.config(cursor=""))
                self.urls_display.tag_bind(tag_name, "<Button-1>", lambda e, url=match.group(): webbrowser.open(url))
        
    def show_search_help(self) -> None:
        """Show a dialog with FTS5 search syntax help."""
        help_text = """Search Tips:
• Use simple keywords to search all fields.
• Use "quotes" for exact phrases: "machine learning".
• Use AND/OR/NOT operators: python AND tutorial.
• Use wildcards: web* (matches web, website, etc.).
• Search specific columns: Purpose:refactor OR Prompt:code.
• Leave empty to show all records."""
        messagebox.showinfo("Search Help", help_text)
        
    def new_item(self) -> None:
        """Open the prompt editing window to create a new prompt."""
        self.open_prompt_window('new')
        
    def duplicate_item(self) -> None:
        """Duplicate the currently selected prompt."""
        if self.current_item:
            self.open_prompt_window('duplicate', self.current_item)
            
    def change_item(self) -> None:
        """Enter in-place editing mode for the selected prompt."""
        if self.current_item and not self.editing_mode:
            self.enter_editing_mode()
            
    def delete_items(self) -> None:
        """Delete one or more selected prompts."""
        if not self.selected_items: return
            
        count = len(self.selected_items)
        if messagebox.askyesno("Confirm Delete", f"Delete {count} item(s)? This cannot be undone."):
            try:
                with self.get_db_connection() as conn:
                    item_ids = tuple(self.selected_items)
                    conn.execute(f"DELETE FROM prompts WHERE id IN ({','.join('?' for _ in item_ids)})", item_ids)
                    for item_id in item_ids:
                        self.clear_prompt_cache(item_id)
                    conn.commit()
                
                self.current_item = None
                self.selected_items = []
                self.clear_item_display()
                self.perform_search(select_first=True)
                self.logger.info(f"Deleted {count} items")
            except Exception as e:
                self.logger.error(f"Delete error: {e}")
                messagebox.showerror("Delete Error", f"Failed to delete: {e}")
                
    def tune_with_ai(self) -> None:
        """Open the AI tuning window for the current prompt."""
        if self.current_item:
            self.open_ai_tuning_window(self.current_item)

    def update_action_buttons(self) -> None:
        """Centralized state machine for managing action buttons."""
        # Hide all buttons first
        for btn in [self.new_btn, self.duplicate_btn, self.change_btn, self.delete_btn, self.save_btn, self.cancel_btn]:
            btn.pack_forget()

        if self.editing_mode:
            # Edit mode: Show Save and Cancel
            self.save_btn.pack(side=tk.LEFT, padx=(0, 5))
            self.cancel_btn.pack(side=tk.LEFT, padx=5)
        else:
            # View mode: Show standard actions
            self.new_btn.pack(side=tk.LEFT, padx=(0, 5))
            
            num_selected = len(self.selected_items)
            
            # Duplicate and Change are only for single selections
            self.duplicate_btn.config(state='normal' if num_selected == 1 else 'disabled')
            self.change_btn.config(state='normal' if num_selected == 1 else 'disabled')
            self.duplicate_btn.pack(side=tk.LEFT, padx=5)
            self.change_btn.pack(side=tk.LEFT, padx=5)
            
            # Delete button state
            self.delete_btn.config(state='normal' if num_selected > 0 else 'disabled')
            self.delete_btn.config(text=f"Delete ({num_selected})" if num_selected > 1 else "Delete")
            self.delete_btn.pack(side=tk.LEFT, padx=(20, 0))

    def save_edits(self) -> None:
        """Save in-place edits to the database."""
        if not self.editing_mode or not self.current_item: return

        try:
            purpose = self.purpose_entry.get()
            prompt = self.prompt_display.get(1.0, tk.END).strip()
            session_urls = self.urls_display.get(1.0, tk.END).strip()
            tags_input = self.tags_entry.get()
            note = self.note_display.get(1.0, tk.END).strip()
            item_id = self.current_item

            tags_json = json.dumps([t.strip() for t in tags_input.split(',') if t.strip()]) if tags_input else None

            with self.get_db_connection() as conn:
                conn.execute('''
                    UPDATE prompts SET Modified = ?, Purpose = ?, Prompt = ?, SessionURLs = ?, Tags = ?, Note = ?
                    WHERE id = ?
                ''', (datetime.now().isoformat(), purpose, prompt, session_urls, tags_json, note, item_id))
                self.clear_prompt_cache(item_id)
                conn.commit()

            self.exit_editing_mode()
            self.perform_search(select_item_id=item_id)
            self.logger.info(f"Saved changes for item {item_id}")
        except Exception as e:
            self.logger.error(f"Error saving edits: {e}")
            messagebox.showerror("Save Error", f"Failed to save changes: {e}")

    def cancel_edits(self) -> None:
        """Cancel in-place editing and restore original content."""
        if self.editing_mode:
            self.exit_editing_mode()
            self.update_item_display(force_refresh=True)

    def enter_editing_mode(self) -> None:
        """Switch the UI to in-place editing mode."""
        if not self.current_item or self.editing_mode: return
            
        try:
            with self.get_db_connection() as conn:
                row = conn.execute('SELECT * FROM prompts WHERE id = ?', (self.current_item,)).fetchone()
                if not row: return

            self.editing_mode = True
            self.update_action_buttons()

            for widget in [self.prompt_display, self.urls_display, self.note_display]:
                widget.config(state='normal')

            self.purpose_display.pack_forget()
            self.purpose_entry = ttk.Entry(self.purpose_frame, font=('TkDefaultFont', 9, 'bold'))
            self.purpose_entry.insert(0, row['Purpose'] or "")
            self.purpose_entry.pack(side=tk.LEFT, padx=(5, 0), fill=tk.X, expand=True)

            for widget in self.tags_display.winfo_children():
                widget.destroy()
            self.tags_entry = ttk.Entry(self.tags_display)
            if row['Tags']:
                try:
                    tags = json.loads(row['Tags']) if row['Tags'].strip().startswith('[') else row['Tags'].split(',')
                    self.tags_entry.insert(0, ', '.join(t.strip() for t in tags))
                except json.JSONDecodeError:
                    self.tags_entry.insert(0, row['Tags']) # fallback
            self.tags_entry.pack(fill=tk.X)
            
        except Exception as e:
            self.logger.error(f"Error entering editing mode: {e}")
            messagebox.showerror("Edit Error", f"Failed to enter editing mode: {e}")
            self.exit_editing_mode() # Rollback UI changes
            
    def exit_editing_mode(self) -> None:
        """Exit editing mode and restore the view-only UI."""
        if not self.editing_mode: return
            
        self.editing_mode = False
        self.update_action_buttons()
        
        for widget in [self.prompt_display, self.urls_display, self.note_display]:
            widget.config(state='disabled')

        if hasattr(self, 'purpose_entry'):
            self.purpose_entry.destroy()
            delattr(self, 'purpose_entry')
        self.purpose_display.pack(side=tk.LEFT, padx=(5, 0))

        if hasattr(self, 'tags_entry'):
            self.tags_entry.destroy()
            delattr(self, 'tags_entry')

    def open_prompt_window(self, mode: str, item_id: Optional[int] = None) -> None:
        """Open a separate window for creating or editing a prompt."""
        window = tk.Toplevel(self.root)
        window.title(f"{mode.title()} Prompt")
        window.transient(self.root)
        window.grab_set()
        window.withdraw()
        
        data = None
        if item_id and mode in ['duplicate', 'change']:
            with self.get_db_connection() as conn:
                data = conn.execute('SELECT * FROM prompts WHERE id = ?', (item_id,)).fetchone()
            
        self.create_prompt_form(window, mode, item_id, data)
        self.root.after(10, lambda: self.auto_size_window(window, 1000, 900, True))
        
    def create_prompt_form(self, window: tk.Toplevel, mode: str, item_id: Optional[int], data: Optional[sqlite3.Row]) -> None:
        """Create the UI components for the prompt editing form."""
        now_str = datetime.now().strftime('%Y-%m-%d %I:%M %p')
        created_text = f"Created: {self.format_datetime(data['Created'])}" if data and mode == 'change' else f"Created: {now_str}"

        date_frame = ttk.Frame(window)
        date_frame.pack(fill=tk.X, padx=10, pady=5)
        ttk.Label(date_frame, text=created_text, foreground="green").pack(side=tk.LEFT)
        ttk.Label(date_frame, text=f"Modified: {now_str}", foreground="blue").pack(side=tk.RIGHT)
        
        purpose_frame = ttk.Frame(window)
        purpose_frame.pack(fill=tk.X, padx=10, pady=5)
        ttk.Label(purpose_frame, text="Purpose:").pack(side=tk.LEFT)
        purpose_var = tk.StringVar(value=data['Purpose'] if data else "")
        purpose_entry = ttk.Entry(purpose_frame, textvariable=purpose_var, width=80)
        purpose_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(5, 0))
        
        prompt_frame = ttk.Frame(window)
        prompt_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        line_frame = ttk.Frame(prompt_frame)
        line_frame.pack(side=tk.LEFT, fill=tk.Y)
        line_numbers = tk.Text(line_frame, width=4, padx=3, takefocus=0, border=0, state='disabled', wrap='none')
        line_numbers.pack(fill=tk.Y, expand=True)
        text_frame = ttk.Frame(prompt_frame)
        text_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)
        prompt_text = tk.Text(text_frame, wrap=tk.WORD, undo=True, maxundo=50)
        text_scrollbar = ttk.Scrollbar(text_frame, orient=tk.VERTICAL)
        prompt_text.config(yscrollcommand=lambda *a: self.sync_scroll(text_scrollbar, line_numbers, *a))
        text_scrollbar.config(command=lambda *a: self.sync_scroll_command(prompt_text, line_numbers, *a))
        prompt_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        text_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        if data and data['Prompt']: prompt_text.insert(1.0, data['Prompt'])
            
        status_frame = ttk.Frame(window)
        status_frame.pack(fill=tk.X, padx=10, pady=5)
        status_label = ttk.Label(status_frame, text="")
        status_label.pack(side=tk.LEFT)
        ttk.Button(status_frame, text="Copy", command=lambda: self.copy_text_to_clipboard(prompt_text.get(1.0, tk.END))).pack(side=tk.RIGHT, padx=(5, 0))
        ttk.Button(status_frame, text="Tune with AI", command=lambda: self.tune_text_with_ai(prompt_text)).pack(side=tk.RIGHT)
        
        def update_form_status(*args: Any) -> None:
            text = prompt_text.get(1.0, tk.END)
            self.update_form_line_numbers(line_numbers, text)
            self.update_form_status_label(status_label, text)
        prompt_text.bind('<KeyRelease>', update_form_status)
        update_form_status()
        
        urls_frame = ttk.LabelFrame(window, text="Session URLs")
        urls_frame.pack(fill=tk.X, padx=10, pady=5)
        urls_text = scrolledtext.ScrolledText(urls_frame, height=3, undo=True, maxundo=50)
        urls_text.pack(fill=tk.X, padx=5, pady=5)
        if data and data['SessionURLs']: urls_text.insert(1.0, data['SessionURLs'])
                
        tags_frame = ttk.LabelFrame(window, text="Tags")
        tags_frame.pack(fill=tk.X, padx=10, pady=5)
        tags_str = ""
        if data and data['Tags']:
            try:
                tags = json.loads(data['Tags']) if data['Tags'].strip().startswith('[') else data['Tags'].split(',')
                tags_str = ', '.join(t.strip() for t in tags)
            except json.JSONDecodeError: tags_str = data['Tags']
        tags_var = tk.StringVar(value=tags_str)
        tags_entry = ttk.Entry(tags_frame, textvariable=tags_var)
        tags_entry.pack(fill=tk.X, padx=5, pady=5)
        
        if WORDCLOUD_AVAILABLE:
            suggestions_frame = ttk.Frame(tags_frame)
            suggestions_frame.pack(fill=tk.X, padx=5, pady=2)
            self.generate_tag_suggestions(suggestions_frame, tags_var, prompt_text)
            
        note_frame = ttk.LabelFrame(window, text="Note")
        note_frame.pack(fill=tk.X, padx=10, pady=5)
        note_text = scrolledtext.ScrolledText(note_frame, height=7, undo=True, maxundo=50)
        note_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        if data and data['Note']: note_text.insert(1.0, data['Note'])
            
        ttk.Button(window, text="Save", command=lambda: self.save_prompt(
            window, mode, item_id, purpose_var.get(), prompt_text.get(1.0, tk.END).strip(),
            urls_text.get(1.0, tk.END).strip(), tags_var.get(), note_text.get(1.0, tk.END).strip()
        )).pack(pady=10)
        
    def update_form_line_numbers(self, line_numbers: tk.Text, text: str) -> None:
        """Update line numbers in a form window."""
        line_numbers.config(state='normal')
        line_numbers.delete(1.0, tk.END)
        if text:
            line_count = text.count('\n') + 1
            line_nums = '\n'.join(map(str, range(1, line_count + 1)))
            line_numbers.insert(1.0, line_nums)
        line_numbers.config(state='disabled')
        
    def update_form_status_label(self, status_label: ttk.Label, text: str) -> None:
        """Update status label in a form window with text statistics."""
        stats = self._get_text_statistics(text)
        status_label.config(
            text=f"Char: {stats.char_count} | Word: {stats.word_count} | "
                 f"Sentence: {stats.sentence_count} | Line: {stats.line_count} | Tokens: {stats.token_count}"
        )
        
    def copy_text_to_clipboard(self, text: str) -> None:
        """Copy arbitrary text to the clipboard."""
        self.root.clipboard_clear()
        self.root.clipboard_append(text)
        self.update_status_bar("Text copied to clipboard")
        
    def tune_text_with_ai(self, text_widget: tk.Text) -> None:
        """Open AI tuning window with text from a given widget."""
        text = text_widget.get(1.0, tk.END).strip()
        if text:
            self.open_ai_tuning_window_with_text(text, text_widget)
            
    def generate_tag_suggestions(self, parent: ttk.Frame, tags_var: tk.StringVar, prompt_text: tk.Text) -> None:
        """Generate keyword-based tag suggestions from the prompt text."""
        if not WORDCLOUD_AVAILABLE: return
            
        def update_suggestions() -> None:
            try:
                text = prompt_text.get(1.0, tk.END).strip()
                if not text: return
                
                words = re.findall(r'\b\w{3,}\b', text.lower())
                common_words = {'the', 'and', 'for', 'with', 'this', 'that', 'are', 'was'}
                word_freq = Counter(w for w in words if w not in common_words)
                
                for widget in parent.winfo_children(): widget.destroy()
                
                for word, _ in word_freq.most_common(7):
                    btn = ttk.Button(parent, text=word, command=lambda w=word: self.add_tag_suggestion(tags_var, w))
                    btn.pack(side=tk.LEFT, padx=2, pady=2)
            except Exception as e:
                self.logger.error(f"Tag suggestion error: {e}")
                
        # Debounce the update
        def on_key_release(event: tk.Event) -> None:
            if self.text_debounce_timer: self.root.after_cancel(self.text_debounce_timer)
            self.text_debounce_timer = self.root.after(1000, update_suggestions)
        
        prompt_text.bind('<KeyRelease>', on_key_release)
        update_suggestions()
        
    def add_tag_suggestion(self, tags_var: tk.StringVar, word: str) -> None:
        """Add a suggested word to the tags entry."""
        current_tags = {t.strip() for t in tags_var.get().split(',') if t.strip()}
        current_tags.add(word)
        tags_var.set(', '.join(sorted(current_tags)))
        
    def save_prompt(self, window: tk.Toplevel, mode: str, item_id: Optional[int], purpose: str, prompt: str, 
                    session_urls: str, tags: str, note: str) -> None:
        """Save a new or updated prompt to the database."""
        try:
            now = datetime.now().isoformat()
            tag_list = [t.strip() for t in tags.split(',') if t.strip()]
            tags_json = json.dumps(tag_list) if tag_list else None
            
            with self.get_db_connection() as conn:
                if mode in ('new', 'duplicate'):
                    cursor = conn.execute('''
                        INSERT INTO prompts (Created, Modified, Purpose, Prompt, SessionURLs, Tags, Note)
                        VALUES (?, ?, ?, ?, ?, ?, ?)
                    ''', (now, now, purpose, prompt, session_urls, tags_json, note))
                    item_id = cursor.lastrowid
                elif mode == 'change' and item_id:
                    conn.execute('''
                        UPDATE prompts SET Modified = ?, Purpose = ?, Prompt = ?, SessionURLs = ?, Tags = ?, Note = ?
                        WHERE id = ?
                    ''', (now, purpose, prompt, session_urls, tags_json, note, item_id))
                    self.clear_prompt_cache(item_id)
                conn.commit()

            window.destroy()
            if item_id:
                self.perform_search(select_item_id=item_id)
            else:
                self.perform_search(select_first=True)
            self.logger.info(f"Saved prompt (mode: {mode})")
            
        except Exception as e:
            self.logger.error(f"Save prompt error: {e}")
            messagebox.showerror("Save Error", f"Failed to save: {e}")
    
    def clear_prompt_cache(self, item_id: Optional[int] = None) -> None:
        """Clear the entire prompt cache or a specific item."""
        if item_id:
            if item_id in self.prompt_cache:
                del self.prompt_cache[item_id]
                self.logger.info(f"Cleared cache for item {item_id}")
        else:
            self.prompt_cache.clear()
            self.logger.info("Cleared entire prompt cache")
            
    def open_ai_tuning_window(self, item_id: int) -> None:
        """Open the AI tuning window for an existing prompt."""
        try:
            with self.get_db_connection() as conn:
                prompt = conn.execute('SELECT Prompt FROM prompts WHERE id = ?', (item_id,)).fetchone()
                if prompt and prompt['Prompt']:
                    self.open_ai_tuning_window_with_text(prompt['Prompt'])
        except Exception as e:
            self.logger.error(f"AI tuning error: {e}")
            messagebox.showerror("AI Tuning Error", f"Failed to open AI tuning: {e}")
            
    def open_ai_tuning_window_with_text(self, text: str, target_widget: Optional[tk.Text] = None) -> None:
        """Open the AI tuning window with pre-filled text."""
        window = tk.Toplevel(self.root)
        window.title("Tune with AI")
        window.transient(self.root)
        window.grab_set()
        window.withdraw()
        
        settings_frame = ttk.LabelFrame(window, text="AI Settings")
        settings_frame.pack(fill=tk.X, padx=10, pady=5)
        provider_frame = ttk.Frame(settings_frame)
        provider_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Label(provider_frame, text="AI Provider:").pack(side=tk.LEFT)
        provider_var = tk.StringVar(value=self.settings_manager.get('ai_provider', 'OpenAI'))
        provider_combo = ttk.Combobox(provider_frame, textvariable=provider_var, 
                                     values=list(AIManager._get_default_settings().keys()),
                                     state="readonly", width=15)
        provider_combo.pack(side=tk.LEFT, padx=(5, 10))
        
        ttk.Label(provider_frame, text="API Key:").pack(side=tk.LEFT)
        api_key_var = tk.StringVar(value=self.settings_manager.get('ai_api_key', ''))
        api_key_entry = ttk.Entry(provider_frame, textvariable=api_key_var, show="*", width=20)
        api_key_entry.pack(side=tk.LEFT, padx=(5, 10))
        
        ttk.Button(provider_frame, text="Get API Key", 
                   command=lambda: self.open_api_key_url(provider_var.get())).pack(side=tk.LEFT, padx=(5, 10))
        
        ttk.Label(provider_frame, text="Model:").pack(side=tk.LEFT)
        model_var = tk.StringVar()
        model_combo = ttk.Combobox(provider_frame, textvariable=model_var, width=25)
        model_combo.pack(side=tk.LEFT, padx=(5, 5))
        
        ttk.Button(provider_frame, text="✏", width=3,
                   command=lambda: self.edit_models(provider_var.get(), model_var)).pack(side=tk.LEFT, padx=(5, 10))
        
        def on_provider_change(*args: Any) -> None:
            provider = provider_var.get()
            provider_defaults = AIManager._get_default_settings().get(provider, {})
            custom_models = self.settings_manager.get('custom_models', {}).get(provider)

            if custom_models:
                model_combo['values'] = custom_models
                model_var.set(custom_models[0] if custom_models else "")
            else:
                models_list = provider_defaults.get('MODELS_LIST', [provider_defaults.get('MODEL', '')])
                model_combo['values'] = [m for m in models_list if m]
                model_var.set(provider_defaults.get('MODEL', ''))
                
        provider_var.trace_add('write', on_provider_change)
        on_provider_change()
        
        main_frame = ttk.Frame(window)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Input and Output panels
        panels = {}
        for side in ['Input', 'Output']:
            frame = ttk.LabelFrame(main_frame, text=side)
            frame.pack(side=tk.LEFT if side == 'Input' else tk.RIGHT, fill=tk.BOTH, expand=True, padx=5)
            
            if side == 'Output':
                btn_frame = ttk.Frame(frame)
                btn_frame.pack(fill=tk.X, padx=5, pady=2)
                ttk.Button(btn_frame, text="Copy to Clipboard", command=lambda: self.copy_text_to_clipboard(panels['Output']['text'].get(1.0, tk.END))).pack(side=tk.LEFT)

            line_frame = ttk.Frame(frame)
            line_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
            line_numbers = tk.Text(line_frame, width=4, padx=3, takefocus=0, border=0, state='disabled', wrap='none')
            line_numbers.pack(side=tk.LEFT, fill=tk.Y)
            
            text_frame = ttk.Frame(line_frame)
            text_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)
            text_widget = tk.Text(text_frame, wrap=tk.WORD, undo=True, maxundo=50, state=('disabled' if side == 'Output' else 'normal'))
            scrollbar = ttk.Scrollbar(text_frame, orient=tk.VERTICAL)
            text_widget.config(yscrollcommand=lambda *a, s=scrollbar, l=line_numbers: self.sync_scroll(s, l, *a))
            scrollbar.config(command=lambda *a, t=text_widget, l=line_numbers: self.sync_scroll_command(t, l, *a))
            text_widget.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
            scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
            
            status_frame = ttk.Frame(frame)
            status_frame.pack(fill=tk.X, padx=5, pady=2)
            status_label = ttk.Label(status_frame, text="")
            status_label.pack(side=tk.LEFT)

            panels[side] = {'text': text_widget, 'lines': line_numbers, 'status': status_label}

        panels['Input']['text'].insert(1.0, f"Please help me improve this AI prompt:\n\n{text}")
        
        ttk.Button(provider_frame, text="Generate AI Response", command=lambda: self.generate_ai_response_with_settings(
            panels['Input']['text'], panels['Output']['text'], panels['Output']['lines'], panels['Output']['status'],
            provider_var.get(), api_key_var.get(), model_var.get()
        )).pack(side=tk.LEFT, padx=(5, 0))
        
        control_frame = ttk.Frame(window)
        control_frame.pack(fill=tk.X, padx=10, pady=5)
        if target_widget:
            ttk.Button(control_frame, text="Apply to Original", command=lambda: self.apply_ai_result(
                panels['Output']['text'], target_widget, window
            )).pack(side=tk.LEFT, padx=5)
        ttk.Button(control_frame, text="Close", command=window.destroy).pack(side=tk.RIGHT, padx=5)
        
        def update_statuses(*args: Any) -> None:
            for side in ['Input', 'Output']:
                t = panels[side]['text'].get(1.0, tk.END)
                self.update_form_line_numbers(panels[side]['lines'], t)
                self.update_form_status_label(panels[side]['status'], t)
                
        panels['Input']['text'].bind('<KeyRelease>', update_statuses)
        update_statuses()
        
        self.root.after(10, lambda: self.auto_size_window(window, 1400, 900, True))
        
    def generate_ai_response_with_settings(self, input_text: tk.Text, output_text: tk.Text, output_lines: tk.Text, 
                                           output_status: ttk.Label, provider: str, api_key: str, model: str) -> None:
        """Generate an AI response using the specified settings in a background thread."""
        input_prompt = input_text.get(1.0, tk.END).strip()
        if not input_prompt: return messagebox.showwarning("No Input", "Please enter text to process")
        if not api_key: return messagebox.showerror("AI Error", "Please enter an API key")
            
        self.settings_manager.set('ai_provider', provider)
        self.settings_manager.set('ai_api_key', api_key)
            
        output_text.config(state='normal')
        output_text.delete(1.0, tk.END)
        output_text.insert(1.0, "Generating AI response...")
        output_text.config(state='disabled')
        
        def ai_worker() -> None:
            try:
                ai_manager = AIManager(tool_name=provider, api_key=api_key)
                override_settings = {'MODEL': model} if model else {}
                response = ai_manager.generate_response(input_prompt, override_settings)
                
                def update_ui() -> None:
                    output_text.config(state='normal')
                    output_text.delete(1.0, tk.END)
                    output_text.insert(1.0, response)
                    output_text.config(state='disabled')
                    self.update_form_line_numbers(output_lines, response)
                    self.update_form_status_label(output_status, response)
                    
                self.root.after(0, update_ui)
            except Exception as e:
                self.logger.error(f"AI generation error: {e}")
                self.root.after(0, lambda: output_text.config(state='normal') or output_text.delete(1.0, tk.END) or output_text.insert(1.0, f"AI Error: {e}") or output_text.config(state='disabled'))
                
        threading.Thread(target=ai_worker, daemon=True).start()
        
    def open_api_key_url(self, provider: str) -> None:
        """Open the appropriate URL for obtaining an API key for the given provider."""
        urls = {
            "Google AI": "https://makersuite.google.com/app/apikey", "Anthropic AI": "https://console.anthropic.com/account/keys",
            "OpenAI": "https://platform.openai.com/api-keys", "Cohere AI": "https://dashboard.cohere.ai/api-keys",
            "HuggingFace AI": "https://huggingface.co/settings/tokens", "Groq AI": "https://console.groq.com/keys",
            "OpenRouterAI": "https://openrouter.ai/keys"
        }
        if provider in urls:
            webbrowser.open(urls[provider])
        else:
            messagebox.showinfo("API Key", f"Please visit the {provider} website for your API key.")
            
    def edit_models(self, provider: str, model_var: tk.StringVar) -> None:
        """Open a dialog to edit the list of available models for a provider."""
        provider_defaults = AIManager._get_default_settings().get(provider, {})
        custom_models = self.settings_manager.get('custom_models', {}).get(provider)
        
        if custom_models:
            models_list = custom_models
        else:
            models_list = provider_defaults.get('MODELS_LIST', [provider_defaults.get('MODEL', '')])
            models_list = [m for m in models_list if m]

        dialog = tk.Toplevel(self.root)
        dialog.title(f"Edit Models - {provider}")
        dialog.transient(self.root)
        dialog.grab_set()
        dialog.withdraw()
        
        ttk.Label(dialog, text="Available Models (one per line, first is default):").pack(pady=5)
        models_text = scrolledtext.ScrolledText(dialog, height=15)
        models_text.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        models_text.insert(1.0, '\n'.join(models_list))
        
        def save_models() -> None:
            new_models = [m.strip() for m in models_text.get(1.0, tk.END).strip().split('\n') if m.strip()]
            if new_models:
                all_custom_models = self.settings_manager.get('custom_models', {})
                all_custom_models[provider] = new_models
                self.settings_manager.set('custom_models', all_custom_models) # This also saves
                
                # Manually trigger the update in the AI window
                provider_combo = window.nametowidget(provider_var.get())
                on_provider_change()

                dialog.destroy()

        btn_frame = ttk.Frame(dialog)
        btn_frame.pack(fill=tk.X, padx=10, pady=5)
        ttk.Button(btn_frame, text="Save", command=save_models).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="Cancel", command=dialog.destroy).pack(side=tk.LEFT, padx=5)
        
        self.root.after(10, lambda: self.auto_size_window(dialog, 500, 400, True))
        
    def apply_ai_result(self, output_text: tk.Text, target_widget: tk.Text, window: tk.Toplevel) -> None:
        """Apply the AI-generated text back to the original text widget."""
        result = output_text.get(1.0, tk.END).strip()
        if result and "Generating AI response..." not in result and "AI Error:" not in result:
            target_widget.delete(1.0, tk.END)
            target_widget.insert(1.0, result)
            window.destroy()
            messagebox.showinfo("Applied", "AI result applied successfully.")
            
    def set_export_location(self) -> None:
        """Open a dialog to set the default export location."""
        folder = filedialog.askdirectory(initialdir=self.settings_manager.get('export_path'))
        if folder:
            self.settings_manager.set('export_path', folder)
            messagebox.showinfo("Export Location", f"Export location set to: {folder}")
            
    def export_view(self, format_type: str) -> None:
        """Export the currently visible search results to a file."""
        if not hasattr(self, 'search_results') or not self.search_results:
            return messagebox.showwarning("No Data", "No items to export.")
        self._export_data(self.search_results, format_type, "view")
            
    def export_all(self, format_type: str) -> None:
        """Export all prompts from the database to a file."""
        try:
            with self.get_db_connection() as conn:
                all_results = conn.execute('SELECT * FROM prompts ORDER BY Modified DESC').fetchall()
            if not all_results:
                return messagebox.showwarning("No Data", "Database is empty.")
            self._export_data(all_results, format_type, "all")
        except Exception as e:
            self.logger.error(f"Export All error: {e}")
            messagebox.showerror("Export Error", f"Failed to fetch data for export: {e}")

    def _export_data(self, data: List[sqlite3.Row], format_type: str, scope: str) -> None:
        """Generic data export handler."""
        exporters = {'csv': self.export_to_csv, 'pdf': self.export_to_pdf, 'txt': self.export_to_txt, 'docx': self.export_to_docx}
        if format_type not in exporters:
            return messagebox.showerror("Export Error", f"Unsupported format: {format_type}")

        try:
            filename = f"prompt_mini_{scope}_{datetime.now():%Y%m%d_%H%M%S}.{format_type}"
            filepath = os.path.join(self.settings_manager.get('export_path'), filename)
            exporters[format_type](data, filepath)
            messagebox.showinfo("Export Complete", f"Exported to: {filepath}")
        except Exception as e:
            self.logger.error(f"Export error (format: {format_type}): {e}")
            messagebox.showerror("Export Error", f"Export failed: {e}")
    
    def export_to_csv(self, data: List[sqlite3.Row], filepath: str) -> None:
        """Export data to a CSV file."""
        if not PANDAS_AVAILABLE: raise ImportError("pandas library is required for CSV export.")
        
        records = []
        for row in data:
            tags = row['Tags']
            tags_str = ""
            if tags:
                try: tags_str = "; ".join(json.loads(tags))
                except (json.JSONDecodeError, TypeError): tags_str = tags
            records.append({
                'ID': row['id'], 'Created': row['Created'], 'Modified': row['Modified'], 'Purpose': row['Purpose'],
                'Prompt': row['Prompt'], 'Session URLs': row['SessionURLs'], 'Tags': tags_str, 'Note': row['Note']
            })
        pd.DataFrame(records).to_csv(filepath, index=False)
        
    def export_to_pdf(self, data: List[sqlite3.Row], filepath: str) -> None:
        """Export data to a PDF document."""
        if not REPORTLAB_AVAILABLE: raise ImportError("reportlab is required for PDF export.")
        
        doc = SimpleDocTemplate(filepath, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        for row in data:
            story.append(Paragraph(f"<b>ID: {row['id']}</b> ({row['Purpose'] or 'No Purpose'})", styles['h2']))
            story.append(Paragraph(f"<i>Created: {self.format_datetime(row['Created'])} | Modified: {self.format_datetime(row['Modified'])}</i>", styles['Normal']))
            
            for field in ['Prompt', 'SessionURLs', 'Note']:
                if row[field]:
                    story.append(Paragraph(f"<b>{field}:</b>", styles['Normal']))
                    story.append(Paragraph(row[field].replace('\n', '<br/>'), styles['BodyText']))
            
            if row['Tags']:
                try: tags_str = ", ".join(json.loads(row['Tags']))
                except (json.JSONDecodeError, TypeError): tags_str = row['Tags']
                story.append(Paragraph(f"<b>Tags:</b> {tags_str}", styles['Normal']))
                
            story.append(Spacer(1, 20))
            
        doc.build(story)
        
    def export_to_txt(self, data: List[sqlite3.Row], filepath: str) -> None:
        """Export data to a plain text file."""
        with open(filepath, 'w', encoding='utf-8') as f:
            for row in data:
                f.write(f"ID: {row['id']}\nPurpose: {row['Purpose'] or ''}\n")
                f.write(f"Created: {self.format_datetime(row['Created'])} | Modified: {self.format_datetime(row['Modified'])}\n")
                
                if row['Tags']:
                    try: tags_str = ", ".join(json.loads(row['Tags']))
                    except (json.JSONDecodeError, TypeError): tags_str = row['Tags']
                    f.write(f"Tags: {tags_str}\n")

                for field in ['Prompt', 'SessionURLs', 'Note']:
                    if row[field]: f.write(f"\n--- {field.upper()} ---\n{row[field]}\n")
                    
                f.write("\n" + "="*80 + "\n\n")
                
    def export_to_docx(self, data: List[sqlite3.Row], filepath: str) -> None:
        """Export data to a DOCX document."""
        if not DOCX_AVAILABLE: raise ImportError("python-docx is required for DOCX export.")
        
        doc = Document()
        doc.add_heading('Prompt Mini Export', 0)
        
        for row in data:
            doc.add_heading(f"ID: {row['id']} - {row['Purpose'] or 'No Purpose'}", level=2)
            doc.add_paragraph(f"Created: {self.format_datetime(row['Created'])} | Modified: {self.format_datetime(row['Modified'])}")
            
            if row['Tags']:
                try: tags_str = ", ".join(json.loads(row['Tags']))
                except (json.JSONDecodeError, TypeError): tags_str = row['Tags']
                p = doc.add_paragraph(); p.add_run('Tags: ').bold = True; p.add_run(tags_str)

            for field in ['Prompt', 'SessionURLs', 'Note']:
                if row[field]:
                    p = doc.add_paragraph(); p.add_run(f'{field}:').bold = True
                    doc.add_paragraph(row[field])
            doc.add_page_break()
        doc.save(filepath)
        
    def backup_database(self) -> None:
        """Create a backup copy of the database file."""
        try:
            with self.get_db_connection() as conn:
                count = conn.execute('SELECT COUNT(*) FROM prompts').fetchone()[0]
            if count == 0: return messagebox.showinfo("No Data", "Database is empty, nothing to backup.")
                
            filename = f"prompt_mini_backup_{datetime.now():%Y%m%d_%H%M%S}.bck"
            backup_path = os.path.join(self.settings_manager.get('export_path'), filename)
            
            shutil.copy2('prompt_mini.db', backup_path)
            messagebox.showinfo("Backup Complete", f"Backup created: {backup_path}")
            self.logger.info(f"Database backed up to {backup_path}")
        except Exception as e:
            self.logger.error(f"Backup error: {e}")
            messagebox.showerror("Backup Error", f"Backup failed: {e}")
            
    def restore_database(self) -> None:
        """Restore the database from a backup file, replacing current data."""
        backup_file = filedialog.askopenfilename(title="Select Backup File", filetypes=[("Backup files", "*.bck")], initialdir=self.settings_manager.get('export_path'))
        if not backup_file: return
            
        if messagebox.askyesno("Confirm Restore", "This will ERASE all current data and replace it with the backup. This cannot be undone. Are you sure?"):
            try:
                # Ensure db is closed by using context manager for a quick op
                with self.get_db_connection() as conn: pass

                shutil.copy2(backup_file, 'prompt_mini.db')
                self.init_database()
                self.perform_search(select_first=True)
                messagebox.showinfo("Restore Complete", "Database restored successfully.")
                self.logger.info(f"Database restored from {backup_file}")
            except Exception as e:
                self.logger.error(f"Restore error: {e}")
                messagebox.showerror("Restore Error", f"Restore failed: {e}")
                
    def import_database(self) -> None:
        """Import records from a backup file into the current database."""
        backup_file = filedialog.askopenfilename(title="Select Backup File to Import", filetypes=[("Backup files", "*.bck")], initialdir=self.settings_manager.get('export_path'))
        if not backup_file: return
            
        try:
            import_conn = sqlite3.connect(backup_file)
            import_conn.row_factory = sqlite3.Row
            import_records = import_conn.execute('SELECT * FROM prompts ORDER BY id').fetchall()
            import_conn.close()
            
            if not import_records: return messagebox.showinfo("No Data", "The backup file is empty.")
            
            duplicate_count = self.analyze_duplicates(import_records)
            
            if self.show_import_confirmation(len(import_records), duplicate_count):
                self.perform_import(import_records)
        except Exception as e:
            self.logger.error(f"Import error: {e}")
            messagebox.showerror("Import Error", f"Failed to read backup file: {e}")
            
    def analyze_duplicates(self, import_records: List[sqlite3.Row]) -> int:
        """Analyze potential duplicates between import records and the current database."""
        try:
            with self.get_db_connection() as conn:
                existing_records = conn.execute('SELECT Purpose, Prompt, Note FROM prompts').fetchall()
            
            existing_set = {(r['Purpose'], r['Prompt'], r['Note']) for r in existing_records}
            import_set = {(r['Purpose'], r['Prompt'], r['Note']) for r in import_records}
            
            return len(existing_set.intersection(import_set))
        except Exception as e:
            self.logger.error(f"Duplicate analysis error: {e}")
            return 0 # Fail safe
        
    def show_import_confirmation(self, total: int, duplicates: int) -> bool:
        """Show a confirmation dialog for importing records, warning about duplicates."""
        dialog = tk.Toplevel(self.root)
        dialog.title("Confirm Import")
        dialog.transient(self.root)
        dialog.grab_set()
        dialog.withdraw()
        
        main_frame = ttk.Frame(dialog)
        main_frame.pack(padx=20, pady=20)
        
        ttk.Label(main_frame, text="⚠️ Import Confirmation", font=('TkDefaultFont', 12, 'bold')).pack(pady=(0, 15))
        
        details_text = (f"Records to import: {total}\n"
                        f"Potential duplicates: {duplicates}\n"
                        f"New unique records: {total - duplicates}")
        ttk.Label(main_frame, text=details_text).pack(pady=(0, 15))
        
        warning_text = "This will add all records from the backup with new IDs.\nDuplicates will not be skipped. This action cannot be undone."
        ttk.Label(main_frame, text=warning_text, wraplength=400, justify=tk.LEFT).pack(pady=(0, 15))
        
        result = {'confirmed': False}
        def confirm() -> None: result['confirmed'] = True; dialog.destroy()
        def cancel() -> None: dialog.destroy()
        
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(fill=tk.X)
        ttk.Button(btn_frame, text="Import Anyway", command=confirm).pack(side=tk.RIGHT)
        ttk.Button(btn_frame, text="Cancel", command=cancel).pack(side=tk.RIGHT, padx=5)
        
        self.root.after(10, lambda: self.auto_size_window(dialog, 450, 250, True))
        dialog.wait_window()
        return result['confirmed']
        
    def perform_import(self, import_records: List[sqlite3.Row]) -> None:
        """Execute the import process, adding records to the database."""
        try:
            with self.get_db_connection() as conn:
                now = datetime.now().isoformat()
                for record in import_records:
                    conn.execute('''
                        INSERT INTO prompts (Created, Modified, Purpose, Prompt, SessionURLs, Tags, Note)
                        VALUES (?, ?, ?, ?, ?, ?, ?)
                    ''', (now, now, record['Purpose'], record['Prompt'], record['SessionURLs'], record['Tags'], record['Note']))
                conn.commit()
            
            self.perform_search(select_first=True)
            messagebox.showinfo("Import Complete", f"Successfully imported {len(import_records)} records.")
            self.logger.info(f"Imported {len(import_records)} records from backup.")
        except Exception as e:
            self.logger.error(f"Import execution error: {e}")
            messagebox.showerror("Import Error", f"Import failed during database write: {e}")
                
    def show_console_log(self) -> None:
        """Show a window with filterable application logs."""
        log_window = tk.Toplevel(self.root)
        log_window.title("Console Log")
        log_window.transient(self.root)
        log_window.withdraw()
        
        level_frame = ttk.Frame(log_window)
        level_frame.pack(fill=tk.X, padx=10, pady=5)
        
        ttk.Label(level_frame, text="Log Level:").pack(side=tk.LEFT)
        level_var = tk.StringVar(value=self.settings_manager.get('log_level', 'INFO'))
        level_combo = ttk.Combobox(level_frame, textvariable=level_var, 
                                  values=["DEBUG", "INFO", "WARNING", "ERROR", "CRITICAL"],
                                  state="readonly", width=10)
        level_combo.pack(side=tk.LEFT, padx=(5, 0))
        
        log_text = scrolledtext.ScrolledText(log_window, state='disabled')
        log_text.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        def update_log_display() -> None:
            level_map = {"DEBUG": 10, "INFO": 20, "WARNING": 30, "ERROR": 40, "CRITICAL": 50}
            selected_level = level_map.get(level_var.get(), 20)
            
            log_text.config(state='normal')
            log_text.delete(1.0, tk.END)
            
            filtered_logs = [msg for lvl, msg in self.log_messages if lvl >= selected_level]
            log_text.insert(tk.END, '\n'.join(filtered_logs))
                    
            log_text.config(state='disabled')
            log_text.see(tk.END)
            
        def on_level_change(event: tk.Event) -> None:
            self.settings_manager.set('log_level', level_var.get())
            self.apply_log_level()
            update_log_display()
            
        level_combo.bind('<<ComboboxSelected>>', on_level_change)
        update_log_display()
        
        def auto_refresh() -> None:
            if log_window.winfo_exists():
                update_log_display()
                log_window.after(2000, auto_refresh)
        auto_refresh()
        
        self.root.after(10, lambda: self.auto_size_window(log_window, 800, 600, True))
        
    def on_closing(self) -> None:
        """Handle application closing events, like saving window geometry."""
        try:
            self.settings_manager.set('window_geometry', self.root.geometry())
            self.logger.info("Saved window geometry and settings.")
        except Exception as e:
            self.logger.error(f"Error saving window geometry: {e}")
        finally:
            self.search_executor.shutdown(wait=False)
            self.root.destroy()
    
    def run(self) -> None:
        """Start the main application event loop."""
        self.root.mainloop()

if __name__ == "__main__":
    app = PromptMiniApp()
    app.run()
